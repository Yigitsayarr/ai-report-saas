from fastapi import FastAPI, Request, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from dotenv import load_dotenv
from openai import OpenAI
import os
import uuid

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX
from docx import Document
from docx.shared import Pt

# =========================
# Ortam & Font
# =========================
load_dotenv()
pdfmetrics.registerFont(TTFont("DejaVu", "DejaVuSans.ttf"))

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = FastAPI(title="AI Rapor + Teklif Oluşturucu")

# =========================
# Static & Template
# =========================
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# =========================
# UI
# =========================
@app.get("/ui", response_class=HTMLResponse)
def ui(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# =========================
# Sağlık
# =========================
@app.get("/")
def root():
    return {"status": "API çalışıyor 🚀"}

# =========================
# ORTAK AI METNİ
# =========================
def generate_ai_text(company_name, sector, problem, requested_service):
    prompt = f"""
Sen deneyimli bir danışmansın.
Aşağıdaki bilgilere dayanarak, doğrudan müşteriye gönderilebilecek
kurumsal teklif hazırla.

FORMAT KURALLARI
Sadece düz metin kullan
Markdown, yıldız, tire, numara, emoji kullanma
Başlıklar BÜYÜK HARF
Kurumsal ve resmi dil kullan

En üstte başlık olmasın direk Sayın, ile başlasın.

FİRMA BİLGİLERİ
Firma Adı: {company_name}
Sektör: {sector}
Problem: {problem}
Talep Edilen Hizmet: {requested_service}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# =========================
# PDF ÜRET & İNDİR
# =========================
@app.post("/generate-pdf")
def generate_pdf(
    report_title: str = Form(...),
    company_name: str = Form(...),
    prepared_by: str = Form(...),
    target_unit: str = Form(...),
    report_date: str = Form(...),
    report_version: str = Form(...),
    sector: str = Form(...),
    problem: str = Form(...),
    requested_service: str = Form(...)
):

    ai_text = generate_ai_text(company_name, sector, problem, requested_service)

    file_name = f"rapor_{uuid.uuid4()}.pdf"
    file_path = os.path.join(os.getcwd(), file_name)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    styles["Title"].fontName = "DejaVu"
    styles["Title"].fontSize = 16
    styles["Title"].spaceAfter = 16

    styles.add(ParagraphStyle(
        name="Meta",
        fontName="DejaVu",
        fontSize=10,
        spaceAfter=4
    ))

    styles.add(ParagraphStyle(
        name="Body",
        fontName="DejaVu",
        fontSize=11,
        leading=16,
        spaceAfter=10
    ))

    elements = []

    elements.append(Paragraph(report_title, styles["Title"]))
    elements.append(Paragraph(f"Kurum / Şirket: {company_name}", styles["Meta"]))
    elements.append(Paragraph(f"Hazırlayan: {prepared_by}", styles["Meta"]))
    elements.append(Paragraph(f"Hitap Edilen Birim: {target_unit}", styles["Meta"]))
    elements.append(Paragraph(f"Tarih: {report_date}", styles["Meta"]))
    elements.append(Paragraph(f"Rapor No / Versiyon: {report_version}", styles["Meta"]))
    elements.append(Spacer(1, 30))

    for line in ai_text.split("\n"):
        if line.strip():
            elements.append(Paragraph(line, styles["Body"]))
        else:
            elements.append(Spacer(1, 10))

    doc.build(elements)

    return FileResponse(
        path=file_path,
        media_type="application/pdf",
        filename=file_name
    )

# =========================
# DOCX ÜRET & İNDİR
# =========================
@app.post("/generate-docx")
def generate_docx(
    report_title: str = Form(...),
    company_name: str = Form(...),
    prepared_by: str = Form(...),
    target_unit: str = Form(...),
    report_date: str = Form(...),
    report_version: str = Form(...),
    sector: str = Form(...),
    problem: str = Form(...),
    requested_service: str = Form(...)
):

    ai_text = generate_ai_text(company_name, sector, problem, requested_service)

    file_name = f"rapor_{uuid.uuid4()}.docx"
    file_path = os.path.join(os.getcwd(), file_name)

    doc = Document()

    title = doc.add_heading(report_title, level=1)
    title.runs[0].font.size = Pt(16)

    meta_fields = [
        f"Kurum / Şirket: {company_name}",
        f"Hazırlayan: {prepared_by}",
        f"Hitap Edilen Birim: {target_unit}",
        f"Tarih: {report_date}",
        f"Rapor No / Versiyon: {report_version}",
    ]

    for m in meta_fields:
        p = doc.add_paragraph()
        run = p.add_run(m)
        run.font.size = Pt(10)

    doc.add_paragraph("")

    for line in ai_text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.save(file_path)

    return FileResponse(
        path=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name
    )
